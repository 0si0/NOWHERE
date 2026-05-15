from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path("/Users/ryu/Project/NOWHERE")
OUT = ROOT / "NEXT Contest 보고서_류현우_NOWHERE.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(Cm(width_cm).twips)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="C8CDD6", sz="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), sz)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text, bold=False, size=10, color="111827"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def apply_document_style(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.space_after = Pt(3)

    for name, size, color, before, after in [
        ("Title", 15, "111827", 0, 8),
        ("Heading 1", 12, "17324D", 9, 4),
        ("Heading 2", 10.5, "17324D", 6, 2),
        ("Heading 3", 10, "374151", 4, 1),
    ]:
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_run(paragraph, text, bold=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(10)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    return run


def add_paragraph(doc, text="", style=None, align=None):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(3)
    if text:
        add_run(p, text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style=None)
    p.style = doc.styles["Normal"]
    p.paragraph_format.left_indent = Cm(0.45)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "· ")
    add_run(p, text)
    return p


def add_section_title(doc, text):
    add_paragraph(doc, text, style="Heading 1")


def add_subtitle(doc, text):
    add_paragraph(doc, text, style="Heading 2")


def add_info_table(doc):
    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    widths = [3.1, 4.0, 3.1, 6.4]
    rows = [
        ("지원형태", "□ 개인   ☑ 팀", "접수번호", "181594-4274"),
        ("신청(대표)자명", "류현우", "팀명", "NOWHERE"),
        ("전화번호", "010-2923-0518", "E-Mail", "nicky0705@naver.com"),
        ("구분", "대표자", "학번", "2024320088"),
        ("학과/학부 / 성함 / 역할", "컴퓨터학과 / 류현우 / 기획자·개발자", "", ""),
    ]
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.rows[r].cells[c]
            set_cell_width(cell, widths[c])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if c in (0, 2):
                set_cell_shading(cell, "EEF2F7")
                set_cell_text(cell, value, bold=True, size=9.5)
            else:
                set_cell_text(cell, value, size=9.5)
    add_paragraph(doc)


def add_summary_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    headers = ["핵심 가치", "구현 근거", "사용자 효과"]
    widths = [4.6, 6.4, 5.6]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, "E8EEF5")
        set_cell_text(cell, h, bold=True, size=9.3)
    rows = [
        ("공간 기반 음악 경험", "장소 반경, 현재 위치, 날씨, 시간대, 청취 기록을 하나의 맥락으로 묶어 추천·기록·공유 기능에 활용", "음악을 단순 재생이 아니라 특정 장소와 기억에 연결"),
        ("현실 제약을 반영한 완성도", "Spotify 개발 모드 5명 제한, Premium 요건, 백그라운드 재생 제약을 반영해 알림 기반 UX와 데모 모드 병행", "심사 환경에서도 주요 흐름을 안정적으로 체험 가능"),
        ("기록과 추천의 순환", "청취 이벤트를 Firestore에 저장하고 추천 슬롯, 뮤직지도, 뮤직다이어리, Shall We Share에 재사용", "쓸수록 개인화가 강화되는 서비스 구조"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_width(cells[i], widths[i])
            set_cell_text(cells[i], value, size=8.7)
    add_paragraph(doc)


def add_feature_table(doc):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    headers = ["기능", "구현 상태", "기술 요소", "기획 대비 조정"]
    widths = [3.1, 3.0, 6.2, 4.3]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, "E8EEF5")
        set_cell_text(cell, h, bold=True, size=8.7)
    rows = [
        ("장소 알림", "완료", "Expo Location/TaskManager, geofence 후보 계산, Firestore 저장 장소, 로컬 쿨다운, 알림 딥링크", "무단 자동재생 대신 도착 알림 후 Spotify 실행"),
        ("추천 탭", "완료", "청취 이벤트, 시간·날씨·장소 컨텍스트, 한국 상황별 후보, Spotify 차트 fallback, OpenAI Challenge", "초기 빈 추천 문제를 fallback으로 해결"),
        ("뮤직지도", "완료", "Kakao Maps WebView, routePoints/segments, 앨범 색상 추출, Spotify Now Playing/순차 데모 모드", "Spotify 권한 한계를 일반/데모 모드로 분리"),
        ("뮤직다이어리", "완료", "지도 경로·앨범아트·청취 시간 기반 영수증형 카드, 저장/공유 흐름", "기록의 재방문성과 공유성을 강화"),
        ("Shall We Share", "완료", "현재 위치 350m 주변 기록 조회, 1일 1회 작성 제한, 익명 장소 핀, 선택 곡 재생", "실시간 위치 노출 대신 사용자가 남긴 장소 음악으로 전환"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_width(cells[i], widths[i])
            set_cell_text(cells[i], value, size=8.2)
    add_paragraph(doc)


def add_architecture_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    headers = ["영역", "주요 파일/모듈", "역할"]
    widths = [3.2, 5.8, 7.6]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, "E8EEF5")
        set_cell_text(cell, h, bold=True, size=8.8)
    rows = [
        ("앱 셸", "App.js, AppNavigator, Session/Location/Player Context", "인증, 위치, 음악 상태를 전역으로 공급하고 홈·추천·지도·다이어리·공유 화면을 연결"),
        ("음악 연동", "nowhere-player, musicPlayerService, Spotify App Remote SDK", "Spotify 앱을 통한 재생 제어, 현재 곡 상태, 큐, 권한 요청, 네이티브 브리지 담당"),
        ("위치/지도", "LocationContext, KakaoMusicMap, KakaoPlacePicker", "권한 요청, 위치 캐시, 지오펜스 후보 탐색, 지도 표시와 장소 선택 처리"),
        ("데이터/백엔드", "firebaseService, firestore.rules, functions/index.js", "Firestore CRUD, 보안 규칙 검증, Spotify/OpenAI 프록시, 사용량 제한과 서버 비밀키 보호"),
        ("개인화", "listeningHistoryService, recommendationService", "시간·날씨·장소 컨텍스트와 청취 이벤트를 기반으로 추천 슬롯을 갱신"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_width(cells[i], widths[i])
            set_cell_text(cells[i], value, size=8.3)
    add_paragraph(doc)


def add_copyright_tables(doc):
    add_section_title(doc, "4. 저작권")
    add_paragraph(doc, "프로젝트에 오픈소스 모듈(라이브러리, 프레임워크, SDK, AI 모델, 외부 API 등)을 사용한 경우, 사용한 모듈명과 적용 영역, 라이선스를 아래와 같이 정리하였다.")
    add_subtitle(doc, "오픈소스 모듈")
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    widths = [2.5, 4.5, 7.2, 2.4]
    headers = ["모듈 유형", "모듈명", "사용한 기능", "라이선스"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, "E8EEF5")
        set_cell_text(cell, h, bold=True, size=8.4)
    modules = [
        ("프레임워크", "React 19.2.0, React Native 0.83.6, Expo 55", "모바일 앱 UI, 런타임, 빌드 및 네이티브 권한 구성", "MIT"),
        ("네비게이션", "@react-navigation/native, bottom-tabs, native-stack", "하단 탭 및 화면 전환 구조", "MIT"),
        ("Expo 모듈", "expo-location, task-manager, media-library, sharing, constants, font, status-bar, dev-client", "위치 권한/백그라운드 작업, 사진 저장, 공유, 설정값, 폰트, 상태바, 개발 빌드", "MIT"),
        ("React Native UI", "@expo/vector-icons, gesture-handler, reanimated, safe-area-context, screens, svg, webview, view-shot, maps, worklets", "아이콘, 제스처, 애니메이션, 안전영역, SVG, WebView 지도, 화면 캡처, 지도 표시", "MIT"),
        ("로컬 저장소", "@react-native-async-storage/async-storage", "위치·날씨·추천·플레이리스트·튜토리얼 캐시", "MIT"),
        ("웹 호환", "react-dom, react-native-web, babel-preset-expo", "Expo Web 및 번들 변환 지원", "MIT"),
        ("Firebase 클라이언트", "firebase 12.11.0", "Authentication, Firestore, Functions 호출", "Apache-2.0"),
        ("Firebase 서버", "firebase-admin 13.8.0, firebase-functions 6.6.0", "Cloud Functions, Firestore Admin, Secret 관리, callable API", "Apache-2.0 / MIT"),
        ("자체 모듈", "nowhere-player", "iOS/Android Spotify 재생 브리지", "프로젝트 자체 코드"),
        ("Spotify SDK", "Spotify iOS SDK / Android App Remote SDK", "Spotify 앱 인증, 재생 제어, 현재 재생곡/PlayerState 연동", "Spotify Developer Terms 및 SDK 조건"),
        ("지도 API", "Kakao Maps JavaScript API", "장소 선택, 음악 지도, 공유 장소 지도 렌더링", "Kakao Developers 이용약관"),
        ("날씨 API", "Open-Meteo Forecast API", "현재 위치의 날씨 코드, 기온, 습도 조회", "Open-Meteo 이용약관"),
        ("AI 모델/API", "OpenAI API, gpt-4.1-nano", "Challenge 추천, 영문 곡명 한국어 표시 보조", "OpenAI 서비스 약관"),
    ]
    for row in modules:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_width(cells[i], widths[i])
            set_cell_text(cells[i], value, size=7.2)

    add_subtitle(doc, "저작권 자료")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    widths = [3.3, 5.5, 7.8]
    headers = ["자료 유형", "자료명", "출처 URL/권리 표기"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, "E8EEF5")
        set_cell_text(cell, h, bold=True, size=8.4)
    resources = [
        ("앱 그래픽", "AppLogo, EmptyMark, ChallengeMark, ChallengeOrb, receipt 이미지", "프로젝트 내부 제작/수정 자산: assets/ 디렉터리"),
        ("아이콘", "Ionicons via @expo/vector-icons", "https://ionic.io/ionicons / MIT"),
        ("음악 메타데이터", "Spotify 곡명, 아티스트명, 앨범아트, URI", "https://developer.spotify.com/ / Spotify API 응답. 음악 음원 파일은 저장·배포하지 않고 Spotify 앱에서 재생"),
        ("지도 타일/장소 표시", "Kakao Maps 지도·타일·장소 UI", "https://apis.map.kakao.com/ / Kakao Maps API 이용"),
        ("날씨 데이터", "Open-Meteo 현재 날씨 데이터", "https://open-meteo.com/ / Forecast API"),
        ("AI 생성 응답", "Challenge 추천 문장 및 곡 추천 결과", "OpenAI API 응답. 최종 재생은 Spotify URI로 연결"),
        ("웹 호스팅 HTML", "hosting/kakao-map.html, music-map.html", "프로젝트 자체 작성 코드, Kakao Maps SDK 로드용"),
    ]
    for row in resources:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_width(cells[i], widths[i])
            set_cell_text(cells[i], value, size=7.3)


def build():
    doc = Document()
    apply_document_style(doc)

    title = add_paragraph(doc, "고려대학교 정보대학 소프트웨어 경진대회 NE:XT Contest 보고서", style="Title", align=WD_ALIGN_PARAGRAPH.CENTER)
    title.paragraph_format.space_after = Pt(8)
    add_info_table(doc)

    add_section_title(doc, "2. 보고서")
    add_subtitle(doc, "1. 프로젝트명")
    add_paragraph(doc, "NOWHERE(나우히어)")

    add_subtitle(doc, "2. 개발 완성도")
    add_paragraph(doc, "NOWHERE는 사용자가 자주 가는 장소, 현재 시간, 날씨, 실제 청취 기록을 연결하여 음악 경험을 공간의 기억으로 확장하는 모바일 서비스다. 단순한 음악 플레이어가 아니라 Spotify 재생을 매개로 장소 알림, 개인화 추천, 뮤직지도, 뮤직다이어리, 익명 장소 공유가 순환하는 ‘나만의 음악 공간’으로 설계하였다.")
    add_summary_table(doc)

    add_paragraph(doc, "프로젝트는 React Native/Expo 기반으로 완성되었고, Firebase Authentication·Firestore·Cloud Functions를 이용해 사용자 데이터, 추천 기록, 공유 기록을 관리한다. Spotify API 정책 변화와 개발 모드 제한은 구현 중 가장 큰 외부 제약이었다. 공식 문서 기준 개발 모드 앱은 허용 사용자 수와 Premium 계정 요건이 있어, 초기 기획의 무조건 백그라운드 자동재생을 그대로 유지하면 심사 환경에서 실패 가능성이 높았다. 따라서 사용자가 장소에 도착하면 NOWHERE가 알림을 띄우고, 사용자가 알림을 누르면 Spotify 앱에서 지정 곡이나 플레이리스트를 여는 구조로 조정하였다. 이 변경은 사용자의 동의 없는 재생을 피하고 플랫폼 정책에도 맞는 방향이다.")
    add_feature_table(doc)

    add_subtitle(doc, "핵심 구현 상세")
    add_paragraph(doc, "장소 알림 기능은 사용자가 지도에서 장소와 반경(50m, 100m, 200m, 300m)을 선택하고 곡 또는 플레이리스트를 연결하는 흐름이다. 저장 가능한 장소는 사용자당 최대 5개로 제한하여 관리 부담을 줄였다. LocationContext는 foreground/background 권한, 위치 캐시, 날씨 캐시, geofence 후보 탐색을 담당하고, autoPlayService는 저장 장소와 현재 위치 사이의 거리를 계산해 반경 진입 여부와 30분 쿨다운을 판단한다. 이 결과는 알림 딥링크로 이어지며, 사용자가 명시적으로 Spotify 실행을 선택한다.")
    add_paragraph(doc, "추천 탭은 ‘요즘 자주 듣는곡’, ‘오늘 이 시간의 추천’, ‘이곳에 어울리는 곡’, ‘오늘 이 날씨의 추천’, ‘Challenge’의 5개 슬롯으로 구성된다. listeningHistoryService가 청취 이벤트와 시간·날씨·장소 컨텍스트를 저장하고, recommendationService는 1순위로 사용자 내부 기록, 2순위로 상황별 한국 음악 후보, 3순위로 Spotify 트렌드/Top 50 fallback을 사용한다. Challenge는 장르·국가·분위기 및 자유 입력을 조합하여 Cloud Functions의 OpenAI 프록시를 통해 새로운 곡을 제안한다.")
    add_paragraph(doc, "뮤직지도는 사용자가 들은 음악을 지도 경로와 앨범아트 색상으로 남기는 기능이다. 일반 모드는 Spotify 현재 재생곡을 가져와 화면이 꺼지거나 일시정지해도 기록이 이어지는 흐름을 목표로 하고, 심사용 제약을 고려한 데모 모드는 사용자가 만든 트랙 플레이리스트의 곡 길이를 기준으로 순차 기록을 진행한다. routePoints, routeSegments, trackChangeMarkers를 기록해 단순 핀 저장이 아니라 음악이 바뀐 구간까지 지도 위에 표현한다.")
    add_paragraph(doc, "뮤직다이어리는 뮤직지도 기록을 다시 꺼내 ‘영수증형 카드’로 정리하는 화면이다. 이동 거리, 재생 시간, 대표 곡, 최대 4개의 요약 트랙, 앨범아트와 경로 미리보기를 제공하여 사용자가 기록을 저장하거나 공유하고 싶게 만드는 후속 경험을 담당한다. Shall We Share는 사용자가 하루 한 번 현재 장소에 음악 한 곡과 80자 이내의 한마디를 남기는 기능이다. 초기 기획의 실시간 주변 사용자 노출은 사생활 위험이 있어, 350m 반경의 익명 장소 기록을 조회하는 방식으로 변경하였다.")

    add_subtitle(doc, "기술적 차별성과 안정성")
    add_architecture_table(doc)
    add_paragraph(doc, "가장 중요한 차별점은 ‘재생 앱을 새로 만드는 것’이 아니라, 기존 Spotify 재생 환경 위에 장소·시간·날씨·기억의 레이어를 얹는 설계다. 음악 파일을 직접 저장하거나 배포하지 않고 Spotify URI와 메타데이터를 사용하며, 사용자의 맥락 데이터만 NOWHERE의 자산으로 축적한다. 또한 API 키와 OpenAI/Spotify owner secret은 Cloud Functions Secret으로 분리하고, 클라이언트에는 공개 가능한 Firebase/Kakao/Spotify Client ID만 둔다.")
    add_paragraph(doc, "Firestore 보안 규칙은 savedPlaces, playRecords, favoriteArtists 등 주요 데이터 구조의 필수 필드, 길이, 좌표 범위, 상태값을 검증한다. 좌표는 위도 -90~90, 경도 -180~180 범위를 검사하고, 장소 반경은 허용 값만 저장된다. 이 구조는 데모 단계에서도 잘못된 데이터가 쌓이는 위험을 줄이며, 추천 품질을 유지하기 위한 데이터 정합성 기반이 된다.")

    add_subtitle(doc, "3. 디자인 및 사용자 경험")
    add_paragraph(doc, "디자인은 어두운 배경(#05070A)에 복숭아색 포인트(#FFC8B8)와 연한 텍스트를 사용해 밤에 음악을 듣는 감각을 유지했다. 홈 화면 중앙의 큰 원은 현재 장소·시간·날씨를 반영한 추천의 중심으로 배치했고, 장소알림설정·뮤직지도·좋아요·이곳에 한마디 기능은 하단 액션 타일로 노출해 한 손 조작이 가능하도록 했다.")
    add_paragraph(doc, "사용자 친화성은 ‘처음 실행해도 비어 있지 않은 앱’에 초점을 맞췄다. 추천은 개인 기록이 없을 때도 상황별 후보와 트렌드 fallback으로 채워지고, SpotlightGuide가 장소 알림·공유 기능의 첫 사용 흐름을 단계적으로 안내한다. 권한이 필요한 기능은 바로 실패시키지 않고 위치 권한, Spotify 권한, Firebase 설정 상태를 확인한 뒤 사용자에게 필요한 행동을 안내한다.")
    add_paragraph(doc, "접근성과 반응형 측면에서는 SafeAreaView, ScrollView, numberOfLines, adjustsFontSizeToFit, compact 레이아웃 분기 등을 사용했다. 긴 곡명은 MarqueeText 또는 줄임 처리로 UI가 깨지지 않게 하고, 지도·카드·필터 영역은 모바일 세로 화면에서 스크롤 가능하게 구성하였다. 음악 지도와 다이어리 화면은 시각적 즐거움을 제공하되, 기본 조작은 선택·기록·저장·공유의 명확한 버튼 중심으로 설계하였다.")

    add_subtitle(doc, "검토 후 반영한 개선 사항")
    add_paragraph(doc, "완성도 점검 과정에서 세 가지 개선 방향을 도출했다. 첫째, Spotify 정책 제약을 보고서에서 명확히 설명해 기획 변경이 미완성이 아니라 합리적 조정임을 드러내야 한다. 둘째, 개인정보 위험이 있는 실시간 위치 공유 대신 사용자가 의도적으로 남긴 장소 기록 중심으로 설명해야 한다. 셋째, 저작권 페이지에는 직접 의존성뿐 아니라 외부 API, SDK, 음악 메타데이터, 지도 타일, 자체 제작 자산을 분리해 기재해야 한다. 본 보고서에는 이 세 가지를 모두 반영했다.")

    doc.add_page_break()
    add_copyright_tables(doc)

    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Cm(1.45)
        section.bottom_margin = Cm(1.3)
        section.left_margin = Cm(1.45)
        section.right_margin = Cm(1.45)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
